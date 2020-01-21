from docx import Document
from docx.shared import Inches
from pg1 import *
from pg2 import *
from pg3 import *
from pg4 import *

document = Document()

document.add_heading(document_title, 0)

for title, body in content_list:
    document.add_heading(title, level= 1)
    document.add_paragraph(body)


document.save('demo.docx')